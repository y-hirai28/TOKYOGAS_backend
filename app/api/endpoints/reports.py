import uuid
import os
from datetime import datetime
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, status, BackgroundTasks
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session

from app.core.security import get_current_user
from app.db.session import get_db
from app.models.user import User
from app.models.report import Report
from app.models.energy_record import EnergyRecord
from app.schemas.report import (
    AutoReportRequest,
    AutoReportPreview,
    AutoReportResponse,
    AutoReportStatus
)

router = APIRouter()

# Directory for generated reports
REPORTS_DIR = "generated_reports"
os.makedirs(REPORTS_DIR, exist_ok=True)


def generate_report_task(report_id: str, request: AutoReportRequest, db_url: str):
    """Background task to generate report"""
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker

    engine = create_engine(db_url, connect_args={"check_same_thread": False})
    SessionLocal = sessionmaker(bind=engine)
    db = SessionLocal()

    try:
        report = db.query(Report).filter(Report.report_id == report_id).first()
        if not report:
            return

        report.progress = 25
        db.commit()

        # Simulate report generation
        import time
        time.sleep(2)

        report.progress = 50
        db.commit()

        # Generate actual file
        file_name = f"report_{report_id}.{request.format}"
        file_path = os.path.join(REPORTS_DIR, file_name)

        if request.format == "pdf":
            # Generate PDF
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas

            c = canvas.Canvas(file_path, pagesize=letter)
            c.drawString(100, 750, f"Energy Report: {request.start_date} to {request.end_date}")
            c.drawString(100, 700, f"Report Type: {request.report_type}")
            c.drawString(100, 650, f"Generated: {datetime.now().isoformat()}")
            c.save()
        else:
            # Generate DOCX
            from docx import Document

            doc = Document()
            doc.add_heading(f"Energy Report", 0)
            doc.add_paragraph(f"Period: {request.start_date} to {request.end_date}")
            doc.add_paragraph(f"Report Type: {request.report_type}")
            doc.add_paragraph(f"Generated: {datetime.now().isoformat()}")
            doc.save(file_path)

        report.progress = 100
        report.status = "completed"
        report.file_name = file_name
        report.file_url = f"/api/v1/reports/generate/download/{report_id}"
        report.completed_at = datetime.utcnow()
        report.size_bytes = os.path.getsize(file_path)
        db.commit()

    except Exception as e:
        report = db.query(Report).filter(Report.report_id == report_id).first()
        if report:
            report.status = "failed"
            report.message = str(e)
            db.commit()
    finally:
        db.close()


@router.post("/generate/preview", response_model=AutoReportPreview)
async def preview_report(
    request: AutoReportRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Preview report before generation"""
    # Count records in date range
    try:
        start = datetime.fromisoformat(request.start_date.replace('Z', '+00:00'))
        end = datetime.fromisoformat(request.end_date.replace('Z', '+00:00'))
    except ValueError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid date format"
        )

    record_count = db.query(EnergyRecord).filter(
        EnergyRecord.timestamp >= start,
        EnergyRecord.timestamp <= end
    ).count()

    sections = ["Executive Summary", "Energy Usage Overview"]
    if request.include_charts:
        sections.extend(["Usage Charts", "Trend Analysis"])
    if request.report_type == "detailed":
        sections.extend(["Detailed Breakdown", "Device Analysis", "Recommendations"])

    estimated_pages = max(1, record_count // 50 + len(sections))

    return AutoReportPreview(
        estimated_pages=estimated_pages,
        sections=sections,
        data_summary={
            "record_count": record_count,
            "date_range": f"{request.start_date} to {request.end_date}",
            "format": request.format
        }
    )


@router.post("/generate", response_model=AutoReportResponse)
async def generate_report(
    request: AutoReportRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Generate a new report"""
    report_id = str(uuid.uuid4())[:8]

    # Create report record
    report = Report(
        report_id=report_id,
        format=request.format,
        report_type=request.report_type,
        status="processing",
        progress=0
    )
    db.add(report)
    db.commit()

    # Start background task
    from app.core.config import settings
    background_tasks.add_task(
        generate_report_task,
        report_id,
        request,
        settings.DATABASE_URL
    )

    return AutoReportResponse(
        report_id=report_id,
        file_url=f"/api/v1/reports/generate/download/{report_id}",
        file_name=f"report_{report_id}.{request.format}",
        format=request.format,
        created_at=report.created_at
    )


@router.get("/generate/status/{report_id}", response_model=AutoReportStatus)
async def get_report_status(
    report_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get report generation status"""
    report = db.query(Report).filter(Report.report_id == report_id).first()

    if not report:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Report not found"
        )

    return AutoReportStatus(
        report_id=report.report_id,
        status=report.status,
        progress=report.progress,
        message=report.message,
        created_at=report.created_at,
        completed_at=report.completed_at
    )


@router.get("/generate/download/{report_id}")
async def download_report(
    report_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Download generated report"""
    report = db.query(Report).filter(Report.report_id == report_id).first()

    if not report:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Report not found"
        )

    if report.status != "completed":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Report not ready for download"
        )

    file_path = os.path.join(REPORTS_DIR, report.file_name)

    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Report file not found"
        )

    media_type = "application/pdf" if report.format == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return FileResponse(
        path=file_path,
        filename=report.file_name,
        media_type=media_type
    )
